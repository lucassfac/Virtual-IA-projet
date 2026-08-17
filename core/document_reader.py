"""
document_reader.py — Lecteur et formateur de documents pour Neural Forge.
"""

import json
import os
from enum import Enum

# ── Seuils ────────────────────────────────────────────────────────────
# Injection directe si contenu ≤ DIRECT_THRESHOLD caractères
DIRECT_THRESHOLD = 12_000   # ~3 000 tokens, tient facilement dans n_ctx=2048

# Taille max d'un chunk pour la résumé en deux passes
CHUNK_SIZE = 8_000          # ~2 000 tokens par chunk

# Limite absolue de caractères injectés dans le prompt final
MAX_INJECT_CHARS = 20_000


class InjectionStrategy(Enum):
    DIRECT  = "direct"   # Contenu brut injecté tel quel
    SUMMARY = "summary"  # Résumé automatique (document trop long)


class DocumentReadError(Exception):
    pass


class DocumentResult:
    """
    Résultat de la lecture d'un document.
    Contient le texte prêt à être injecté et la stratégie utilisée.
    """

    def __init__(
        self,
        content: str,
        strategy: InjectionStrategy,
        original_chars: int,
        filename: str,
    ):
        self.content = content                  # Texte final à injecter
        self.strategy = strategy                # DIRECT ou SUMMARY
        self.original_chars = original_chars    # Taille avant traitement
        self.filename = filename

    @property
    def was_summarized(self) -> bool:
        return self.strategy == InjectionStrategy.SUMMARY

    def build_prompt(self, user_question: str = "") -> str:
        """
        Construit le prompt final à envoyer au LLM.
        """
        if self.was_summarized:
            header = (
                f"Voici un résumé automatique du document « {self.filename} » "
                f"(document original : {self.original_chars:,} caractères) :\n\n"
                f"{self.content}"
            )
        else:
            header = (
                f"Voici le contenu du document « {self.filename} » :\n\n"
                f"{self.content}"
            )

        if user_question:
            return f"{header}\n\n---\nQuestion : {user_question}"
        else:
            return (
                f"{header}\n\n---\n"
                "Fais un résumé clair et structuré de ce document."
            )


# ── Entrée principale ─────────────────────────────────────────────────

def read_document(
    path: str,
    llm_node=None,
    direct_threshold: int = DIRECT_THRESHOLD,
) -> DocumentResult:
    """
    Lit un fichier et décide automatiquement de la stratégie :
      - Court → DocumentResult.strategy = DIRECT
      - Long  → si llm_node fourni : résumé en deux passes (SUMMARY)
               sinon : troncature avec avertissement (DIRECT tronqué)

    :param path:             Chemin du fichier
    :param llm_node:         LLMNode optionnel pour la résumé automatique
    :param direct_threshold: Seuil en caractères pour basculer en résumé
    """
    if not os.path.exists(path):
        raise DocumentReadError(f"Fichier introuvable : {path}")

    filename = os.path.basename(path)
    raw = _extract_text(path)
    original_chars = len(raw)

    if original_chars <= direct_threshold:
        # ── Injection directe ──
        return DocumentResult(
            content=raw,
            strategy=InjectionStrategy.DIRECT,
            original_chars=original_chars,
            filename=filename,
        )

    # ── Document long ──
    if llm_node is not None and llm_node.model_loaded:
        summary = _summarize_in_chunks(raw, llm_node, filename)
        return DocumentResult(
            content=summary,
            strategy=InjectionStrategy.SUMMARY,
            original_chars=original_chars,
            filename=filename,
        )

    # Pas de LLM disponible → troncature intelligente avec avertissement
    truncated = _truncate(raw, MAX_INJECT_CHARS)
    note = (
        f"\n\n[⚠ Document tronqué : {original_chars:,} → {MAX_INJECT_CHARS:,} car. "
        "Chargez un modèle pour activer le résumé automatique.]"
    )
    return DocumentResult(
        content=truncated + note,
        strategy=InjectionStrategy.DIRECT,
        original_chars=original_chars,
        filename=filename,
    )


# ── Résumé en deux passes ─────────────────────────────────────────────

def _summarize_in_chunks(text: str, llm_node, filename: str) -> str:
    """
    Résumé automatique par chunks :
    1. Découpe le texte en chunks de CHUNK_SIZE caractères
    2. Résume chaque chunk séparément
    3. Concatène les résumés partiels
    4. Si le résultat est encore trop long → passe finale de résumé
    """
    from core.types import DataPacket, DataType

    chunks = _split_chunks(text, CHUNK_SIZE)
    partial_summaries = []

    for i, chunk in enumerate(chunks):
        prompt = (
            f"Résume ce passage ({i + 1}/{len(chunks)}) "
            f"du document « {filename} » en 3-5 phrases concises :\n\n{chunk}"
        )
        try:
            packet = DataPacket(DataType.TEXT, prompt, source="document_reader")
            llm_node.set_input(packet)
            result = llm_node._run_inference()
            partial_summaries.append(f"[Partie {i + 1}]\n{result.content.strip()}")
        except Exception:
            # En cas d'échec sur un chunk : on garde le chunk tronqué
            partial_summaries.append(
                f"[Partie {i + 1}]\n{_truncate(chunk, 500)}"
            )

    combined = "\n\n".join(partial_summaries)

    # Passe finale si les résumés partiels sont encore trop longs
    if len(combined) > DIRECT_THRESHOLD and len(chunks) > 1:
        final_prompt = (
            f"Voici les résumés partiels du document « {filename} ». "
            "Synthétise-les en un résumé global cohérent et structuré :\n\n"
            f"{combined}"
        )
        try:
            from core.types import DataPacket, DataType
            packet = DataPacket(DataType.TEXT, final_prompt, source="document_reader")
            llm_node.set_input(packet)
            result = llm_node._run_inference()
            return result.content.strip()
        except Exception:
            pass

    return combined


# ── Extracteurs par format ────────────────────────────────────────────

def _extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md", ".rst", ".csv", ".log"):
        return _read_text(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    elif ext == ".docx":
        return _read_docx(path)
    elif ext in (".json", ".jsonl"):
        return _read_json(path)
    else:
        return _read_text(path)  # Tentative texte brut


def _read_image_fast(path: str) -> str:
    """Extraction instantanée via OCR classique (contourne LLaVA)."""
    try:
        import pytesseract
        from PIL import Image
        
        img = Image.open(path)
        # On force la lecture en français et en anglais
        text = pytesseract.image_to_string(img, lang='fra+eng').strip()
        
        if not text:
            raise DocumentReadError(
                "L'OCR rapide n'a détecté aucun texte. "
                "Le fichier nécessite une analyse visuelle profonde (LLaVA)."
            )
        return text
    except ImportError:
        raise DocumentReadError(
            "Le module 'pytesseract' n'est pas installé. "
            "Tapez : pip install pytesseract"
        )
    except Exception as e:
        raise DocumentReadError(f"Échec de l'OCR rapide : {e}")


def _read_text(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise DocumentReadError(f"Encodage inconnu : '{path}'")


def _read_pdf(path: str) -> str:
    try:
        import fitz
    except ImportError:
        raise DocumentReadError(
            "PyMuPDF non installé — lancez : pip install pymupdf"
        )
    try:
        doc = fitz.open(path)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                pages.append(f"[Page {i + 1}]\n{text}")
        doc.close()
        if not pages:
            raise DocumentReadError(
                "Aucun texte extractible (PDF scanné ou protégé)."
            )
        return "\n\n".join(pages)
    except DocumentReadError:
        raise
    except Exception as e:
        raise DocumentReadError(f"Erreur PDF : {e}")


def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise DocumentReadError(
            "python-docx non installé — lancez : pip install python-docx"
        )
    try:
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception as e:
        raise DocumentReadError(f"Erreur DOCX : {e}")


def _read_json(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".jsonl":
            lines = []
            with open(path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                        lines.append(json.dumps(obj, ensure_ascii=False, indent=2))
                    except json.JSONDecodeError:
                        lines.append(line)
            return "\n---\n".join(lines)
        else:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)
    except Exception as e:
        raise DocumentReadError(f"Erreur JSON : {e}")


# ── Utilitaires ───────────────────────────────────────────────────────

def _split_chunks(text: str, chunk_size: int) -> list[str]:
    """Découpe le texte en chunks en respectant les fins de paragraphes."""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            chunks.append(text[start:])
            break
        # Cherche une coupure propre (fin de paragraphe)
        cut = text.rfind("\n\n", start, end)
        if cut == -1 or cut <= start:
            cut = text.rfind("\n", start, end)
        if cut == -1 or cut <= start:
            cut = end
        chunks.append(text[start:cut])
        start = cut
    return [c.strip() for c in chunks if c.strip()]


def _truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    last_nl = truncated.rfind("\n")
    if last_nl > max_chars * 0.85:
        truncated = truncated[:last_nl]
    return truncated


def get_file_info(path: str) -> dict:
    stat = os.stat(path)
    size_kb = stat.st_size / 1024
    ext = os.path.splitext(path)[1].lower()
    label = {
        ".pdf": "PDF", ".docx": "Word", ".txt": "Texte",
        ".md": "Markdown", ".json": "JSON", ".jsonl": "JSONL",
    }.get(ext, "Fichier")
    return {
        "name": os.path.basename(path),
        "label": label,
        "size_kb": round(size_kb, 1),
        "ext": ext,
    }
